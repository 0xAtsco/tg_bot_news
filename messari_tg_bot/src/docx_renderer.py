import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Literal

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


ItemType = Literal["research", "newsletter"]


@dataclass
class ContentPayload:
    item_id: str
    slug: str
    title: str
    url: str
    publish_date: datetime
    item_type: ItemType
    translated_content: str


class DocxRenderer:
    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def render(self, payload: ContentPayload) -> Path:
        document = Document()

        # Title
        title = document.add_heading(payload.title, level=0)
        title_format = title.runs[0].font
        title_format.size = Pt(18)
        title_format.bold = True

        # Metadata
        meta = document.add_paragraph()
        meta.add_run(f"Опубликовано: ").bold = True
        meta.add_run(payload.publish_date.strftime("%d.%m.%Y %H:%M"))

        meta = document.add_paragraph()
        meta.add_run(f"Тип: ").bold = True
        meta.add_run(payload.item_type.title())

        # URL as clickable link
        url_para = document.add_paragraph()
        url_para.add_run("Оригинал: ").bold = True
        url_run = url_para.add_run(payload.url)
        url_run.font.color.rgb = RGBColor(0, 0, 255)
        url_run.font.underline = True

        # Separator
        document.add_paragraph("_" * 80)

        # Main content with better formatting
        self._add_formatted_content(document, payload.translated_content)

        filename = self._build_filename(payload)
        document.save(filename)
        return filename

    def _add_formatted_content(self, document: Document, content: str) -> None:
        """Add formatted content to document."""
        if not content:
            document.add_paragraph("(пустой текст)")
            return

        paragraphs = self._split_paragraphs(content)

        for para_text in paragraphs:
            if not para_text.strip():
                continue

            # Check if it's a heading (ends with : or is short and capitalized)
            if self._is_heading(para_text):
                heading = document.add_heading(para_text, level=1)
                heading_format = heading.runs[0].font
                heading_format.size = Pt(14)
                heading_format.bold = True
            # Check if it's a list item
            elif para_text.strip().startswith(('- ', '• ', '* ', '– ')):
                p = document.add_paragraph(para_text.strip()[2:], style='List Bullet')
                p.paragraph_format.left_indent = Pt(20)
            # Check if it's a numbered list
            elif re.match(r'^\d+[\.)]\s', para_text.strip()):
                # Remove number prefix and add as numbered list
                text = re.sub(r'^\d+[\.)]\s+', '', para_text.strip())
                p = document.add_paragraph(text, style='List Number')
                p.paragraph_format.left_indent = Pt(20)
            # Check if it's a link line
            elif para_text.strip().startswith(('http://', 'https://', '<http')):
                url = para_text.strip().strip('<>')
                p = document.add_paragraph()
                url_run = p.add_run(url)
                url_run.font.color.rgb = RGBColor(0, 0, 255)
                url_run.font.underline = True
                url_run.font.size = Pt(10)
            # Regular paragraph
            else:
                p = document.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15

    def _is_heading(self, text: str) -> bool:
        """Determine if text should be formatted as a heading."""
        text = text.strip()

        # Ends with colon
        if text.endswith(':'):
            return True

        # Short, all caps (like chapter titles)
        if len(text) < 50 and text.isupper():
            return True

        # Timestamp pattern (00:00)
        if re.match(r'^\d{1,2}:\d{2}', text):
            return True

        return False

    def _build_filename(self, payload: ContentPayload) -> Path:
        prefix = "Messari_Research" if payload.item_type == "research" else "Messari_Newsletter"
        date_part = payload.publish_date.date().isoformat()
        slug_part = payload.slug or payload.item_id
        safe_slug = "".join(ch if ch.isalnum() or ch in "-_" else "-" for ch in slug_part)
        return self.output_dir / f"{prefix}_{date_part}_{safe_slug}.docx"

    @staticmethod
    def _split_paragraphs(content: str) -> list[str]:
        if not content:
            return ["(пустой текст)"]
        parts = [segment.strip() for segment in content.split("\n") if segment.strip()]
        return parts or ["(пустой текст)"]
